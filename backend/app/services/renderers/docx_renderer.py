# =============================================================================
# File: docx_renderer.py
# Module/Service: Report Service (FR9) — DOCX Renderer
# Layer: Service
# Purpose: Render AggregatedReportBlock list into a Word ``.docx`` file (UC8).
# Responsibilities:
#   - Heading styles per block title; body by source_type (mirrors Markdown logic)
#   - Real Word tables for tabular extractions (not screenshots / images)
#   - Write ``{title}_{report_id}.docx`` under report staging path / MinIO key
# Dependencies:
#   - python-docx; report_aggregation.AggregatedReportBlock; renderers.common
# Public Exports:
#   - render_docx, DocxRenderResult
# Database/Table: N/A (file artifact only; reports.file_path set by Report Service)
# Related Modules: markdown_renderer, report_aggregation, app.ai.ocr.docx_parser
# Important Notes:
#   - Reuses project python-docx dependency (no separate docx-helper skill).
#   - Does not alter AggregatedReportBlock schema; no Markdown/PDF here.
# =============================================================================

from __future__ import annotations

import json
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Sequence

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.models.enums import ReportSourceType
from app.services.report_aggregation import AggregatedReportBlock
from app.services.renderers.common import (
    build_report_filename,
    build_report_object_key,
    cell_str,
    ensure_parent_dir,
    extraction_as_table,
    resolve_report_staging_path,
)


@dataclass(frozen=True, slots=True)
class DocxRenderResult:
    """Local staging artifact + MinIO object key for later upload."""

    filename: str
    local_path: Path
    object_key: str
    section_count: int


def render_docx(
    blocks: Sequence[AggregatedReportBlock],
    *,
    report_title: str,
    report_id: uuid.UUID,
    workspace_id: uuid.UUID,
    output_dir: Path | None = None,
) -> DocxRenderResult:
    """Render blocks to DOCX and write the staging ``.docx`` file."""
    filename = build_report_filename(report_title, report_id, extension="docx")
    object_key = build_report_object_key(
        workspace_id=workspace_id,
        report_id=report_id,
        filename=filename,
    )
    local_path = resolve_report_staging_path(
        workspace_id=workspace_id,
        report_id=report_id,
        filename=filename,
        output_dir=output_dir,
    )

    document = Document()
    title = report_title.strip() or "Report"
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for block in blocks:
        document.add_heading(block.title, level=1)
        _append_block_body(document, block)

    ensure_parent_dir(local_path)
    document.save(str(local_path))

    return DocxRenderResult(
        filename=filename,
        local_path=local_path,
        object_key=object_key,
        section_count=len(blocks),
    )


def _append_block_body(document: Document, block: AggregatedReportBlock) -> None:
    source_type = block.source_type
    if isinstance(source_type, str):
        try:
            source_type = ReportSourceType(source_type)
        except ValueError:
            _add_preformatted(document, json.dumps(block.content, ensure_ascii=False, indent=2))
            return

    if source_type is ReportSourceType.summary:
        _append_summary(document, block.content)
    elif source_type is ReportSourceType.comparison:
        _append_comparison(document, block.content)
    elif source_type is ReportSourceType.extraction:
        _append_extraction(document, block.content)
    elif source_type is ReportSourceType.chat_session:
        _append_chat(document, block.content)
    else:
        _add_preformatted(document, json.dumps(block.content, ensure_ascii=False, indent=2))


def _append_summary(document: Document, content: dict[str, Any]) -> None:
    text = content.get("text")
    if isinstance(text, str) and text.strip():
        document.add_paragraph(text.strip())

    sections = content.get("sections")
    if not isinstance(sections, list):
        return
    for section in sections:
        if not isinstance(section, dict):
            continue
        title = cell_str(section.get("title")).strip()
        body = cell_str(section.get("content")).strip()
        if title:
            document.add_heading(title, level=2)
        if body:
            document.add_paragraph(body)


def _append_comparison(document: Document, content: dict[str, Any]) -> None:
    report = content.get("comparison_report")
    if isinstance(report, dict):
        _append_contract_comparison(document, report, content)
        return
    _append_legacy_comparison(document, content)


def _append_legacy_comparison(document: Document, content: dict[str, Any]) -> None:
    similarities = content.get("similarities") or []
    differences = content.get("differences") or []

    if similarities:
        document.add_heading("Similarities", level=2)
        for item in similarities:
            document.add_paragraph(cell_str(item), style="List Bullet")

    if differences:
        document.add_heading("Differences", level=2)
        for item in differences:
            document.add_paragraph(cell_str(item), style="List Bullet")


def _append_contract_comparison(
    document: Document,
    report: dict[str, Any],
    content: dict[str, Any],
) -> None:
    meta = report.get("metadata") if isinstance(report.get("metadata"), dict) else {}
    summary = (
        report.get("executive_summary")
        if isinstance(report.get("executive_summary"), dict)
        else {}
    )
    stats = (
        report.get("overall_statistics")
        if isinstance(report.get("overall_statistics"), dict)
        else {}
    )
    risks = report.get("risk_summary") if isinstance(report.get("risk_summary"), dict) else {}
    generation = (
        report.get("generation_metadata")
        if isinstance(report.get("generation_metadata"), dict)
        else {}
    )

    document.add_heading("Executive Summary", level=2)
    if meta.get("comparison_id"):
        document.add_paragraph(f"Comparison ID: {meta['comparison_id']}")
    if meta.get("generated_at"):
        document.add_paragraph(f"Generated at: {meta['generated_at']}")
    if meta.get("status"):
        document.add_paragraph(f"Status: {meta['status']}")
    for label, key in (
        ("Total clauses", "total_clauses"),
        ("UNCHANGED", "unchanged"),
        ("MODIFIED", "modified"),
        ("ADDED", "added"),
        ("REMOVED", "removed"),
    ):
        document.add_paragraph(f"{label}: {summary.get(key, 0)}", style="List Bullet")
    risk_counts = summary.get("risk_counts") if isinstance(summary.get("risk_counts"), dict) else {}
    document.add_paragraph(
        f"CRITICAL risks: {risk_counts.get('CRITICAL', 0)}",
        style="List Bullet",
    )
    document.add_paragraph(
        f"HIGH risks: {risk_counts.get('HIGH', 0)}",
        style="List Bullet",
    )
    document.add_paragraph(
        f"Verified evidence references: {summary.get('verified_evidence_count', 0)}",
        style="List Bullet",
    )

    document.add_heading("Documents", level=2)
    for doc in report.get("documents") or []:
        if not isinstance(doc, dict):
            continue
        side = cell_str(doc.get("side")) or "Document"
        title = cell_str(doc.get("title")).strip()
        version = cell_str(doc.get("document_version_id")).strip()
        text = side
        if title:
            text += f": {title}"
        if version:
            text += f" (version {version})"
        document.add_paragraph(text, style="List Bullet")

    document.add_heading("Overall Statistics", level=2)
    document.add_paragraph(
        f"Clauses compared: {stats.get('total_clauses_compared', 0)}",
        style="List Bullet",
    )
    if stats.get("verification_rate") is not None:
        document.add_paragraph(
            f"Verification rate: {stats.get('verification_rate')}",
            style="List Bullet",
        )

    document.add_heading("Risk Summary", level=2)
    level_rows = [
        [cell_str(item.get("level")), cell_str(item.get("count"))]
        for item in risks.get("by_level") or []
        if isinstance(item, dict)
    ]
    if level_rows:
        _add_table(document, ["Risk level", "Count"], level_rows)
    category_rows = [
        [cell_str(item.get("category")), cell_str(item.get("count"))]
        for item in risks.get("by_category") or []
        if isinstance(item, dict)
    ]
    if category_rows:
        _add_table(document, ["Risk category", "Count"], category_rows)
    for item in risks.get("items") or []:
        if not isinstance(item, dict):
            continue
        clause = cell_str(item.get("clause_id")) or "Clause"
        level = cell_str(item.get("risk_level")) or "—"
        category = cell_str(item.get("risk_category")) or "—"
        document.add_paragraph(f"{clause} — {level} / {category}", style="List Bullet")
        if item.get("reason"):
            document.add_paragraph(f"Detected risk: {item['reason']}")
        if item.get("explanation"):
            document.add_paragraph(f"Risk explanation: {item['explanation']}")
        if item.get("recommendation"):
            document.add_paragraph(f"Recommendation: {item['recommendation']}")

    _append_clause_table(document, "Changed Clauses", report.get("changed_clauses"))
    _append_clause_table(document, "Added Clauses", report.get("added_clauses"))
    _append_clause_table(document, "Removed Clauses", report.get("removed_clauses"))

    unchanged = (
        report.get("unchanged_clauses")
        if isinstance(report.get("unchanged_clauses"), dict)
        else {}
    )
    document.add_heading("Unchanged Clauses", level=2)
    document.add_paragraph(f"{unchanged.get('count', 0)} clauses remained unchanged.")

    details = report.get("detailed_clause_comparisons") or []
    if details:
        document.add_heading("Detailed Evidence", level=2)
        for detail in details:
            if isinstance(detail, dict):
                _append_clause_detail(document, detail)

    document.add_heading("Generation Metadata", level=2)
    document.add_paragraph(f"Builder: {generation.get('builder') or 'cmp-24'}")
    document.add_paragraph(f"Report LLM calls: {generation.get('llm_calls_report', 0)}")
    document.add_paragraph(f"Upstream LLM calls: {generation.get('llm_calls_upstream', 0)}")
    if generation.get("quality_status"):
        document.add_paragraph(f"Quality status: {generation['quality_status']}")

    if content.get("similarities") or content.get("differences"):
        document.add_heading("Contextual notes", level=2)
        _append_legacy_comparison(document, content)


def _append_clause_table(document: Document, title: str, rows: object) -> None:
    document.add_heading(title, level=2)
    items = [item for item in (rows or []) if isinstance(item, dict)]
    if not items:
        document.add_paragraph("None.")
        return
    _add_table(
        document,
        ["Clause", "Status", "Risk", "Category", "Change"],
        [
            [
                cell_str(item.get("display_id") or item.get("clause_id")),
                cell_str(item.get("status")),
                cell_str(item.get("risk_level")) or "—",
                cell_str(item.get("risk_category")) or "—",
                cell_str(item.get("change")) or "—",
            ]
            for item in items
        ],
    )


def _append_clause_detail(document: Document, detail: dict[str, Any]) -> None:
    display = cell_str(detail.get("display_id") or detail.get("clause_id")) or "Clause"
    document.add_heading(f"Clause {display}", level=3)
    document.add_paragraph(f"Status: {detail.get('status') or '—'}")
    if detail.get("risk_level"):
        document.add_paragraph(f"Risk: {detail['risk_level']}")
    if detail.get("risk_category"):
        document.add_paragraph(f"Category: {detail['risk_category']}")
    if detail.get("verification_status"):
        document.add_paragraph(f"Citation verification: {detail['verification_status']}")
    if detail.get("v1_text") is not None:
        document.add_paragraph("V1").runs[0].bold = True
        document.add_paragraph(str(detail["v1_text"]))
    if detail.get("v2_text") is not None:
        document.add_paragraph("V2").runs[0].bold = True
        document.add_paragraph(str(detail["v2_text"]))
    diffs = detail.get("exact_differences") or []
    if diffs:
        document.add_paragraph("Exact Changes").runs[0].bold = True
        for item in diffs:
            if not isinstance(item, dict):
                continue
            document.add_paragraph(cell_str(item.get("label")) or "Value", style="List Bullet")
            if item.get("old") is not None:
                document.add_paragraph(f"V1: {item['old']}")
            if item.get("new") is not None:
                document.add_paragraph(f"V2: {item['new']}")
            if item.get("delta"):
                document.add_paragraph(f"Absolute delta: {item['delta']}")
            if item.get("percent"):
                document.add_paragraph(f"Percentage change: {item['percent']}")
    if detail.get("explanation"):
        document.add_paragraph("Risk Explanation").runs[0].bold = True
        document.add_paragraph(cell_str(detail["explanation"]))
    if detail.get("recommendation"):
        document.add_paragraph("Recommendation").runs[0].bold = True
        document.add_paragraph(cell_str(detail["recommendation"]))
    if detail.get("absence_note"):
        document.add_paragraph("Absence / counterpart").runs[0].bold = True
        document.add_paragraph(cell_str(detail["absence_note"]))
    evidence = detail.get("evidence") or []
    if not evidence:
        return
    document.add_paragraph("Evidence").runs[0].bold = True
    for item in evidence:
        if not isinstance(item, dict):
            continue
        side = cell_str(item.get("side")) or "Source"
        state = _evidence_state_label(item.get("verification_state"))
        page = item.get("page_number")
        page_bit = f", page {page}" if page not in (None, "") else ""
        document.add_paragraph(f"{side}{page_bit} — {state}", style="List Bullet")
        if item.get("display_text"):
            document.add_paragraph(f"Source: {item['display_text']}")


def _evidence_state_label(state: object) -> str:
    key = cell_str(state).strip().lower()
    if key == "verified":
        return "Verified"
    if key == "partial":
        return "Partially verified"
    if key == "unavailable":
        return "No evidence"
    return "Unverified"


def _append_extraction(document: Document, content: dict[str, Any]) -> None:
    table = extraction_as_table(content)
    if table is not None:
        headers, rows = table
        _add_table(document, headers, rows)
        return

    result = content.get("result")
    payload = result if result is not None else content
    _add_preformatted(document, json.dumps(payload, ensure_ascii=False, indent=2))


def _append_chat(document: Document, content: dict[str, Any]) -> None:
    messages = content.get("messages") or []
    if not isinstance(messages, list):
        _add_preformatted(document, json.dumps(content, ensure_ascii=False, indent=2))
        return

    for message in messages:
        if not isinstance(message, dict):
            continue
        role = cell_str(message.get("role")).strip().lower()
        text = cell_str(message.get("content")).strip()
        if role == "user":
            label = "User"
        elif role == "assistant":
            label = "Assistant"
        else:
            label = role.title() or "Message"
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
        paragraph.add_run(text)


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> Table:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_idx, row in enumerate(rows):
        cells = table.rows[row_idx + 1].cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = value
    return table


def _add_preformatted(document: Document, text: str) -> Paragraph:
    """Monospace-ish paragraph for JSON payloads (no screenshot tables)."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    return paragraph
