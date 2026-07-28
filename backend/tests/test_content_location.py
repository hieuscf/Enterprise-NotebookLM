# =============================================================================
# File: test_content_location.py
# Module/Service: Citation / Chunking location sync (FR5)
# Layer: Service
# Purpose: Verify DOCX section_index vs PDF page_number through OCR→chunking.
# Responsibilities:
#   - DOCX: page_number null, section_index set on TextChunk
#   - PDF: page_number set, section_index null
#   - ContentLocation schema helper does not invent missing fields
# Dependencies:
#   - pytest, app.ai.ocr, app.ai.chunking, app.schemas.content_location
# Public Exports:
#   - N/A
# Database/Table: document_chunks (contract covered via TextChunk fields)
# Related Modules: P0.1 ocr.py; stage_chunking
# Important Notes: Full chat citation HTTP path not implemented yet — schema ready.
# =============================================================================

from __future__ import annotations

import io

import fitz
import pytest
from docx import Document as DocxDocument

from app.ai.chunking import run_chunking_from_segments
from app.ai.ocr import run_ocr_cleaning
from app.models.enums import FileType
from app.schemas.content_location import (
    ContentLocation,
    CitationResponse,
    content_location_from_chunk,
)


def test_docx_ocr_to_chunks_keeps_section_index_not_page() -> None:
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("1. Giới thiệu", level=1)
    doc.add_paragraph("Nội dung mục giới thiệu để citation trích dẫn.")
    doc.add_heading("2. Phương pháp", level=1)
    doc.add_paragraph("Nội dung phương pháp nằm ở mục giữa tài liệu.")
    doc.add_heading("3. Kết luận", level=1)
    doc.add_paragraph("Kết luận ngắn.")
    doc.save(buf)

    ocr = run_ocr_cleaning(file_type=FileType.docx, data=buf.getvalue())
    assert all(s.page_number is None for s in ocr.segments)

    chunks = run_chunking_from_segments(
        [
            {
                "text": s.text,
                "page_number": s.page_number,
                "section": s.section,
                "section_index": s.section_index,
                "order_index": s.order_index,
            }
            for s in ocr.segments
        ]
    )
    assert chunks
    mid = [c for c in chunks if c.section and "Phương pháp" in (c.section or "")]
    assert mid, "expected a chunk under Phương pháp section"
    for c in mid:
        assert c.page_number is None
        assert c.section_index is not None
        assert c.section_index >= 1

    loc = content_location_from_chunk(
        page_number=mid[0].page_number,
        section_index=mid[0].section_index,
        section=mid[0].section,
    )
    assert loc.page_number is None
    assert loc.section_index == mid[0].section_index
    assert loc.section_title == mid[0].section

    # Citation response shape ready for GET .../citations
    citation = CitationResponse(
        id="00000000-0000-4000-8000-000000000001",
        message_id="00000000-0000-4000-8000-000000000002",
        retrieval_id="00000000-0000-4000-8000-000000000003",
        document_id="00000000-0000-4000-8000-000000000004",
        text_snippet=mid[0].content[:200],
        verified=True,
        order_index=0,
        location=loc,
    )
    payload = citation.model_dump()
    assert payload["location"]["page_number"] is None
    assert payload["location"]["section_index"] == mid[0].section_index


def test_pdf_ocr_to_chunks_keeps_page_number_not_section_index() -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "PDF body paragraph for citation page test.", fontsize=11)
    data = doc.tobytes()
    doc.close()

    ocr = run_ocr_cleaning(file_type=FileType.pdf, data=data)
    assert any(s.page_number == 1 for s in ocr.segments)

    chunks = run_chunking_from_segments(
        [
            {
                "text": s.text,
                "page_number": s.page_number,
                "section": s.section,
                "section_index": s.section_index,
                "order_index": s.order_index,
            }
            for s in ocr.segments
        ]
    )
    assert chunks
    for c in chunks:
        assert c.page_number is not None
        assert c.section_index is None

    loc = content_location_from_chunk(
        page_number=chunks[0].page_number,
        section_index=chunks[0].section_index,
        section=chunks[0].section,
    )
    assert loc.page_number == 1
    assert loc.section_index is None


def test_content_location_does_not_invent_fields() -> None:
    loc = ContentLocation(page_number=None, section_index=2, section_title="Intro")
    assert loc.page_number is None
    assert loc.section_index == 2

    loc2 = content_location_from_chunk(page_number=4, section_index=None, section=None)
    assert loc2.page_number == 4
    assert loc2.section_index is None


def test_unmerged_table_candidates_logged_when_block_between() -> None:
    from app.ai.ocr import (
        _ParsedBlock,
        _count_unmerged_table_candidates,
        _format_table_semantic,
    )

    t1 = _format_table_semantic(["A", "B"], [["1", "2"]])
    t2 = _format_table_semantic(["A", "B"], [["3", "4"]])
    blocks = [
        _ParsedBlock(text=t1, page_number=1, section="S", block_type="table", table_col_count=2),
        _ParsedBlock(text="caption between", page_number=2, section="S", block_type="paragraph"),
        _ParsedBlock(text=t2, page_number=2, section="S", block_type="table", table_col_count=2),
    ]
    assert _count_unmerged_table_candidates(blocks) == 1
