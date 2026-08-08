# =============================================================================
# File: test_report_renderers.py
# Module/Service: Report Service (FR9) — Markdown / DOCX / PDF Renderers
# Layer: Service
# Purpose: Unit tests for report Markdown/DOCX/PDF renderers (Prompt 2–3/5).
# Responsibilities:
#   - Output non-empty; section_count == input block count
#   - Filename / object_key conventions; tabular extraction shapes
#   - PDF from Markdown preserves title / headings / table / code text
# Dependencies:
#   - pytest, python-docx, pymupdf, AggregatedReportBlock fakes
# Public Exports:
#   - N/A
# Database/Table: N/A (temp filesystem only; no MinIO/Postgres)
# Related Modules: app.services.renderers
# Important Notes: Does not call LLM or upload to MinIO.
# =============================================================================

from __future__ import annotations

import re
import uuid
from pathlib import Path

import fitz
import pytest
from docx import Document

from app.models.enums import ReportSourceType
from app.services.report_aggregation import AggregatedReportBlock
from app.services.renderers.docx_renderer import render_docx
from app.services.renderers.markdown_renderer import render_markdown
from app.services.renderers.pdf_renderer import markdown_to_html, render_pdf


def _blocks() -> list[AggregatedReportBlock]:
    return [
        AggregatedReportBlock(
            order_index=0,
            source_type=ReportSourceType.summary,
            title="Summary (short) — Policy",
            content={"text": "Policy overview.", "style": "short", "sections": None},
        ),
        AggregatedReportBlock(
            order_index=1,
            source_type=ReportSourceType.extraction,
            title="Extraction (table) — Policy",
            content={
                "extraction_type": "table",
                "result": {
                    "headers": ["Name", "Value"],
                    "rows": [{"Name": "Alpha", "Value": "1"}, {"Name": "Beta", "Value": "2"}],
                },
            },
        ),
        AggregatedReportBlock(
            order_index=2,
            source_type=ReportSourceType.comparison,
            title="Q1 vs Q2",
            content={
                "similarities": ["Same theme"],
                "differences": ["Different scope"],
            },
        ),
        AggregatedReportBlock(
            order_index=3,
            source_type=ReportSourceType.chat_session,
            title="Research chat",
            content={
                "messages": [
                    {"role": "user", "content": "What is policy?", "created_at": "2026-01-01T10:00:00+00:00"},
                    {
                        "role": "assistant",
                        "content": "A set of rules.",
                        "created_at": "2026-01-01T10:01:00+00:00",
                    },
                ]
            },
        ),
        AggregatedReportBlock(
            order_index=4,
            source_type=ReportSourceType.extraction,
            title="Extraction (entities) — Policy",
            content={
                "extraction_type": "entities",
                "result": {"entities": [{"name": "Acme", "type": "ORG"}]},
            },
        ),
    ]


@pytest.fixture
def staging_root(tmp_path: Path) -> Path:
    return tmp_path / "staging"


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


def test_markdown_output_non_empty_and_section_count(staging_root: Path) -> None:
    blocks = _blocks()
    report_id = uuid.uuid4()
    workspace_id = uuid.uuid4()

    result = render_markdown(
        blocks,
        report_title="Quarterly Report",
        report_id=report_id,
        workspace_id=workspace_id,
        output_dir=staging_root,
    )

    assert result.section_count == len(blocks)
    assert result.markdown.strip()
    assert result.local_path.is_file()
    assert result.local_path.stat().st_size > 0
    assert result.filename == f"Quarterly_Report_{report_id}.md"
    assert result.object_key == (
        f"workspaces/{workspace_id}/reports/{report_id}/{result.filename}"
    )

    h2_titles = re.findall(r"^## (.+)$", result.markdown, flags=re.MULTILINE)
    assert len(h2_titles) == len(blocks)
    assert h2_titles == [b.title for b in blocks]

    assert "| Name | Value |" in result.markdown
    assert "**User:** What is policy?" in result.markdown
    assert "**Assistant:** A set of rules." in result.markdown
    assert "```json" in result.markdown  # non-table extraction


def test_markdown_empty_blocks_still_writes_title(staging_root: Path) -> None:
    report_id = uuid.uuid4()
    result = render_markdown(
        [],
        report_title="Empty",
        report_id=report_id,
        workspace_id=uuid.uuid4(),
        output_dir=staging_root,
    )
    assert result.section_count == 0
    assert result.markdown.strip()
    assert result.local_path.read_text(encoding="utf-8").startswith("# Empty")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_docx_output_non_empty_and_section_count(staging_root: Path) -> None:
    blocks = _blocks()
    report_id = uuid.uuid4()
    workspace_id = uuid.uuid4()

    result = render_docx(
        blocks,
        report_title="Quarterly Report",
        report_id=report_id,
        workspace_id=workspace_id,
        output_dir=staging_root,
    )

    assert result.section_count == len(blocks)
    assert result.local_path.is_file()
    assert result.local_path.stat().st_size > 0
    assert result.filename.endswith(".docx")
    assert result.filename == f"Quarterly_Report_{report_id}.docx"
    assert str(report_id) in result.object_key

    document = Document(str(result.local_path))
    heading1 = [
        p.text
        for p in document.paragraphs
        if p.style is not None and p.style.name == "Heading 1"
    ]
    assert len(heading1) == len(blocks)
    assert heading1 == [b.title for b in blocks]

    # Real Word table for tabular extraction (not an image)
    assert len(document.tables) >= 1
    table = document.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["Name", "Value"]
    assert [c.text for c in table.rows[1].cells] == ["Alpha", "1"]

    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "User: What is policy?" in all_text
    assert "Assistant: A set of rules." in all_text


def test_docx_and_markdown_share_section_count(staging_root: Path) -> None:
    blocks = _blocks()
    report_id = uuid.uuid4()
    workspace_id = uuid.uuid4()
    kwargs = {
        "report_title": "Shared",
        "report_id": report_id,
        "workspace_id": workspace_id,
        "output_dir": staging_root,
    }
    md = render_markdown(blocks, **kwargs)
    docx = render_docx(blocks, **kwargs)
    assert md.section_count == docx.section_count == len(blocks)


# ---------------------------------------------------------------------------
# PDF (from Markdown)
# ---------------------------------------------------------------------------


def test_pdf_from_markdown_non_empty_preserves_structure(staging_root: Path) -> None:
    blocks = _blocks()
    report_id = uuid.uuid4()
    workspace_id = uuid.uuid4()
    md = render_markdown(
        blocks,
        report_title="Quarterly Report",
        report_id=report_id,
        workspace_id=workspace_id,
        output_dir=staging_root,
    )

    result = render_pdf(
        md.markdown,
        report_title="Quarterly Report",
        report_id=report_id,
        workspace_id=workspace_id,
        output_dir=staging_root,
    )

    assert result.local_path.is_file()
    assert result.local_path.suffix == ".pdf"
    assert result.local_path.stat().st_size > 0
    assert result.filename == f"Quarterly_Report_{report_id}.pdf"
    assert result.object_key.endswith(result.filename)

    with fitz.open(str(result.local_path)) as doc:
        assert len(doc) >= 1
        text = "\n".join(page.get_text() for page in doc)

    assert "Quarterly Report" in text
    assert "Summary (short)" in text
    assert "Policy overview" in text
    assert "Name" in text and "Alpha" in text
    assert "User:" in text and "What is policy?" in text
    # JSON code-fence content from non-table extraction
    assert "entities" in text or "Acme" in text


def test_markdown_to_html_keeps_headings_table_and_code() -> None:
    md = "\n".join(
        [
            "# Title",
            "",
            "## Section",
            "",
            "- bullet",
            "",
            "| A | B |",
            "| --- | --- |",
            "| 1 | 2 |",
            "",
            "```json",
            '{"k": true}',
            "```",
            "",
            "**User:** hi",
        ]
    )
    html_out = markdown_to_html(md)
    assert "<h1>Title</h1>" in html_out
    assert "<h2>Section</h2>" in html_out
    assert "<li>bullet</li>" in html_out
    assert "<table>" in html_out and "<th>A</th>" in html_out
    assert "<pre><code" in html_out
    assert "<b>User:</b> hi" in html_out
