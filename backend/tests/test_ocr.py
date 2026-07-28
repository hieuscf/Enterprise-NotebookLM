# =============================================================================
# File: test_ocr.py
# Module/Service: Pipeline Worker — OCR & Cleaning ([AI])
# Layer: Service
# Purpose: Unit tests for FR2 Step 3 OCR segments, cleaning, empty-text fail.
# Responsibilities:
#   - TXT/DOCX-like segments; EmptyOcrError; stage metadata + artifact save
# Dependencies:
#   - pytest, app.ai.ocr, app.workers.stages.ocr_cleaning (mocked I/O)
# Public Exports:
#   - N/A
# Database/Table: N/A
# Related Modules: app.ai.ocr, app.workers.artifacts
# Important Notes: No live MinIO/Postgres in CI.
# =============================================================================

from __future__ import annotations

import io
import uuid
from contextlib import contextmanager
from datetime import UTC, datetime
from unittest.mock import MagicMock, patch

import pytest

from app.ai.ocr import EmptyOcrError, run_ocr_cleaning
from app.models.documents import Document, DocumentVersion
from app.models.enums import DocumentVersionStatus, FileType
from app.workers.artifacts import OCR_SEGMENTS_ARTIFACT, pipeline_artifact_key
from app.workers.stages.errors import DataPipelineError
from app.workers.stages.ocr_cleaning import stage_ocr_cleaning


def test_txt_produces_ordered_segments() -> None:
    data = b"First paragraph about NotebookLM.\n\nSecond paragraph about LightRAG.\n"
    result = run_ocr_cleaning(file_type=FileType.txt, data=data)
    assert result.page_count == 1
    assert result.char_count > 0
    assert len(result.segments) >= 2
    assert result.segments[0].order_index == 0
    assert result.segments[1].order_index == 1
    assert "NotebookLM" in result.segments[0].text
    # Legacy adapter for chunking
    assert len(result.pages) == len(result.segments)


def test_empty_text_raises_empty_ocr_error() -> None:
    with pytest.raises(EmptyOcrError, match="scanned PDF|text layer"):
        run_ocr_cleaning(file_type=FileType.txt, data=b"   \n\n\t  ")


def test_whitespace_and_encoding_cleaned() -> None:
    raw = "Hello\u00a0World\ufeff\n\n\nNext   line".encode("utf-8")
    result = run_ocr_cleaning(file_type=FileType.txt, data=raw)
    joined = " ".join(s.text for s in result.segments)
    assert "\ufeff" not in joined
    assert "  " not in joined
    assert "Hello World" in joined


def test_pipeline_artifact_key() -> None:
    key = pipeline_artifact_key(
        "workspaces/ws/documents/doc/v1/report.pdf",
        OCR_SEGMENTS_ARTIFACT,
    )
    assert key == "workspaces/ws/documents/doc/v1/.pipeline/ocr_segments.json"


def test_stage_ocr_cleaning_persists_artifact_and_metadata() -> None:
    version_id = uuid.uuid4()
    doc_id = uuid.uuid4()
    version = DocumentVersion(
        id=version_id,
        document_id=doc_id,
        uploaded_by=uuid.uuid4(),
        version_number=1,
        storage_path="workspaces/ws/documents/doc/v1/a.txt",
        file_size_bytes=20,
        checksum_sha256="x",
        page_count=None,
        status=DocumentVersionStatus.processing,
        is_current=True,
        created_at=datetime.now(UTC),
    )
    document = Document(
        id=doc_id,
        workspace_id=uuid.uuid4(),
        current_version_id=version_id,
        title="A",
        file_type=FileType.txt,
        created_at=datetime.now(UTC),
        updated_at=datetime.now(UTC),
    )

    fake_storage = MagicMock()
    fake_storage.download_bytes.return_value = b"Enterprise NotebookLM segment one.\n\nTwo."
    uploaded: dict[str, bytes] = {}

    def _upload(*, object_key: str, data: bytes, content_type: str = "") -> str:
        uploaded[object_key] = data
        return object_key

    fake_storage.upload_bytes.side_effect = _upload

    @contextmanager
    def _session():
        session = MagicMock()
        session.get.side_effect = lambda model, pk: {
            DocumentVersion: version,
            Document: document,
        }.get(model)
        yield session

    with (
        patch("app.workers.stages.ocr_cleaning.get_minio_storage", return_value=fake_storage),
        patch("app.workers.stages.ocr_cleaning.get_sync_session", _session),
    ):
        meta = stage_ocr_cleaning(version_id)

    assert meta["page_count"] == 1
    assert meta["segment_count"] >= 1
    assert meta["char_count"] > 0
    assert meta["output_bytes"] > 0
    assert "duration_ms" in meta
    assert meta["artifact_key"].endswith(".pipeline/ocr_segments.json")
    assert meta["artifact_key"] in uploaded


def test_stage_ocr_empty_file_fails_with_data_error() -> None:
    version_id = uuid.uuid4()
    doc_id = uuid.uuid4()
    version = DocumentVersion(
        id=version_id,
        document_id=doc_id,
        uploaded_by=uuid.uuid4(),
        version_number=1,
        storage_path="workspaces/ws/documents/doc/v1/empty.txt",
        file_size_bytes=0,
        checksum_sha256="x",
        page_count=None,
        status=DocumentVersionStatus.processing,
        is_current=True,
        created_at=datetime.now(UTC),
    )
    document = Document(
        id=doc_id,
        workspace_id=uuid.uuid4(),
        current_version_id=version_id,
        title="Empty",
        file_type=FileType.txt,
        created_at=datetime.now(UTC),
        updated_at=datetime.now(UTC),
    )

    fake_storage = MagicMock()
    fake_storage.download_bytes.return_value = b"\n\n   "

    @contextmanager
    def _session():
        session = MagicMock()
        session.get.side_effect = lambda model, pk: {
            DocumentVersion: version,
            Document: document,
        }.get(model)
        yield session

    with (
        patch("app.workers.stages.ocr_cleaning.get_minio_storage", return_value=fake_storage),
        patch("app.workers.stages.ocr_cleaning.get_sync_session", _session),
        pytest.raises(DataPipelineError, match="text layer|No extractable text"),
    ):
        stage_ocr_cleaning(version_id)


def test_docx_keeps_heading_as_section() -> None:
    from docx import Document as DocxDocument

    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Body under introduction.")
    doc.save(buf)
    result = run_ocr_cleaning(file_type=FileType.docx, data=buf.getvalue())
    assert any(s.section == "Introduction" for s in result.segments)
    assert any("Body under introduction" in s.text for s in result.segments)


def test_soft_line_breaks_joined_hard_breaks_kept() -> None:
    data = b"Hello\nworld\n\nParagraph A\n\nParagraph B\n"
    result = run_ocr_cleaning(file_type=FileType.txt, data=data)
    texts = [s.text for s in result.segments]
    assert any(t == "Hello world" for t in texts)
    assert any(t == "Paragraph A" for t in texts)
    assert any(t == "Paragraph B" for t in texts)


def test_xlsx_emits_semantic_row_text() -> None:
    from openpyxl import Workbook

    buf = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "People"
    ws.append(["Name", "Age"])
    ws.append(["John", 18])
    wb.save(buf)
    result = run_ocr_cleaning(file_type=FileType.xlsx, data=buf.getvalue())
    joined = "\n".join(s.text for s in result.segments)
    assert "Name = John" in joined
    assert "Age = 18" in joined
    assert " | " not in joined
    assert any(s.block_type == "table" for s in result.segments)
    assert result.metrics.table_count >= 1


def test_docx_table_and_list() -> None:
    from docx import Document as DocxDocument

    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Bullet item one", style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "ColumnA"
    table.cell(0, 1).text = "ColumnB"
    table.cell(1, 0).text = "ValueA"
    table.cell(1, 1).text = "ValueB"
    doc.save(buf)
    result = run_ocr_cleaning(file_type=FileType.docx, data=buf.getvalue())
    joined = "\n".join(s.text for s in result.segments)
    assert "ColumnA = ValueA" in joined
    assert "ColumnB = ValueB" in joined
    assert "Bullet item one" in joined
    assert any(s.section == "Overview" for s in result.segments)


def test_normalize_quotes_dashes_zero_width() -> None:
    raw = "Say \u201cHello\u201d\u2014world\u200b.".encode("utf-8")
    result = run_ocr_cleaning(file_type=FileType.txt, data=raw)
    text = result.segments[0].text
    assert "\u201c" not in text and "\u201d" not in text
    assert "\u2014" not in text and "\u200b" not in text
    assert '"Hello"' in text
    assert "-world" in text


# ---------------------------------------------------------------------------
# P0.1 — DOCX uses section_index, not fake physical page_number
# ---------------------------------------------------------------------------


def test_docx_uses_section_index_not_fake_page_number() -> None:
    """DOCX has no physical pages; page_number must stay None (option 1a).

    Builds a multi-heading document that previously inflated logical_page on
    every heading — which Citation (FR5) would mis-report as \"trang X\".
    """
    from docx import Document as DocxDocument

    buf = io.BytesIO()
    doc = DocxDocument()
    # Preamble before any heading (same physical "page" in Word as heading 1)
    doc.add_paragraph("Preamble text before first heading.")
    doc.add_heading("1. Giới thiệu", level=1)
    for _ in range(8):
        doc.add_paragraph(
            "Đoạn nội dung dài trong phần giới thiệu để mô phỏng nhiều dòng "
            "trên cùng một trang Word vật lý trước khi sang heading tiếp theo."
        )
    # Heading mid-document — must NOT invent a new physical page_number
    doc.add_heading("2. Phương pháp", level=1)
    doc.add_paragraph("Nội dung phương pháp nằm ngay sau heading trên cùng trang logic.")
    doc.add_heading("3. Kết luận", level=1)
    doc.add_paragraph("Kết luận ngắn.")
    doc.save(buf)

    result = run_ocr_cleaning(file_type=FileType.docx, data=buf.getvalue())

    assert all(s.page_number is None for s in result.segments), (
        "DOCX must not invent physical page_number values"
    )
    assert any(s.section_index is not None for s in result.segments)

    by_section_name = {s.section: s.section_index for s in result.segments if s.section}
    assert by_section_name.get("1. Giới thiệu") is not None
    assert by_section_name.get("2. Phương pháp") is not None
    assert by_section_name.get("3. Kết luận") is not None
    assert by_section_name["1. Giới thiệu"] < by_section_name["2. Phương pháp"]
    assert by_section_name["2. Phương pháp"] < by_section_name["3. Kết luận"]

    # page_count for DOCX = logical section count, not Word pagination
    assert result.page_count == max(s.section_index or 0 for s in result.segments)
    # Legacy CleanedPage locator falls back to section_index
    assert all(p.page_number == (s.section_index or s.order_index + 1)
               for p, s in zip(result.pages, result.segments, strict=True))


# ---------------------------------------------------------------------------
# P0.2 — Cross-page PDF table merge
# ---------------------------------------------------------------------------


def test_pdf_merges_table_split_across_pages() -> None:
    from app.ai.ocr import (
        _ParsedBlock,
        _merge_cross_page_tables,
        _format_table_semantic,
    )

    page1 = _format_table_semantic(
        ["Name", "Age"],
        [["John", "18"], ["Jane", "20"]],
    )
    # Continuation reprints header labels as a fake first row (common PDF split)
    page2_with_header_repeat = (
        "Row 1\nName = Name\nAge = Age\nRow 2\nName = Bob\nAge = 22"
    )

    blocks = [
        _ParsedBlock(
            text=page1,
            page_number=1,
            section="Data",
            block_type="table",
            table_col_count=2,
        ),
        _ParsedBlock(
            text=page2_with_header_repeat,
            page_number=2,
            section="Data",
            block_type="table",
            table_col_count=2,
        ),
    ]
    merged = _merge_cross_page_tables(blocks)
    assert len(merged) == 1
    assert merged[0].page_number == 1
    text = merged[0].text
    assert "Name = John" in text
    assert "Name = Jane" in text
    assert "Name = Bob" in text
    # Repeated header labels should not remain as a data row
    assert "Name = Name" not in text
    assert merged[0].table_col_count == 2


def test_pdf_does_not_merge_tables_when_heading_between() -> None:
    from app.ai.ocr import _ParsedBlock, _merge_cross_page_tables, _format_table_semantic

    t1 = _format_table_semantic(["A", "B"], [["1", "2"]])
    t2 = _format_table_semantic(["A", "B"], [["3", "4"]])
    blocks = [
        _ParsedBlock(text=t1, page_number=1, section="S1", block_type="table", table_col_count=2),
        _ParsedBlock(
            text="New Section",
            page_number=2,
            section="New Section",
            block_type="heading",
            heading_level=1,
        ),
        _ParsedBlock(text=t2, page_number=2, section="New Section", block_type="table", table_col_count=2),
    ]
    merged = _merge_cross_page_tables(blocks)
    assert len(merged) == 3


# ---------------------------------------------------------------------------
# P1 — Header/footer must not strip numbered TOC headings
# ---------------------------------------------------------------------------


def test_header_footer_keeps_numbered_heading_at_page_edge() -> None:
    from app.ai.ocr import _ParsedBlock, _strip_repeated_headers_footers

    blocks = [
        _ParsedBlock(
            text="1. Giới thiệu\nNội dung trang một dài hơn.\n12",
            page_number=1,
            block_type="paragraph",
        ),
        _ParsedBlock(
            text="Acme Corp Report\nNội dung trang hai.\n12",
            page_number=2,
            block_type="paragraph",
        ),
        _ParsedBlock(
            text="Acme Corp Report\nNội dung trang ba.\n12",
            page_number=3,
            block_type="paragraph",
        ),
        _ParsedBlock(
            text="Acme Corp Report\nNội dung trang bốn.\n12",
            page_number=4,
            block_type="paragraph",
        ),
    ]
    stripped = _strip_repeated_headers_footers(blocks)
    page1 = next(b for b in stripped if b.page_number == 1)
    assert "1. Giới thiệu" in page1.text

    # Majority header "Acme Corp Report" stripped from pages 2–4; TOC heading kept
    for b in stripped:
        if b.page_number and b.page_number >= 2:
            assert "Acme Corp Report" not in b.text
    assert any("1. Giới thiệu" in b.text for b in stripped)

def test_protected_content_line_helpers() -> None:
    from app.ai.ocr import _is_protected_content_line, _is_boilerplate_line

    assert _is_protected_content_line("1. Giới thiệu")
    assert _is_protected_content_line("2. Phương pháp nghiên cứu")
    assert _is_boilerplate_line("12")
    assert not _is_protected_content_line("12")
    assert _is_boilerplate_line("Page 3")
    assert not _is_protected_content_line("Page 3")


# ---------------------------------------------------------------------------
# P2 — Median sampling performance
# ---------------------------------------------------------------------------


def test_median_sample_matches_full_on_small_input() -> None:
    from app.ai.ocr import _median

    values = [10.0, 11.0, 12.0, 12.0, 100.0]
    assert _median(values) == 12.0


def test_median_sampling_fast_on_large_font_list() -> None:
    import time

    from app.ai.ocr import MEDIAN_SAMPLE_MAX, _median

    # Simulate span sizes from a ≥300-page PDF (many spans/page)
    n = 300 * 80  # 24_000 sizes
    values = [10.0 + (i % 7) * 0.1 for i in range(n)]
    assert len(values) > MEDIAN_SAMPLE_MAX

    t0 = time.perf_counter()
    sampled = _median(values)
    sampled_ms = (time.perf_counter() - t0) * 1000

    t1 = time.perf_counter()
    import statistics

    full = float(statistics.median(values))
    full_ms = (time.perf_counter() - t1) * 1000

    # Sampled median should be close to full median
    assert abs(sampled - full) <= 0.5
    # And not slower than a full sort on this size (usually much faster)
    assert sampled_ms < full_ms * 3 or sampled_ms < 50.0


def test_pdf_median_path_on_multipage_document() -> None:
    """End-to-end: build a 300-page text PDF and ensure OCR completes quickly."""
    import time

    import fitz

    doc = fitz.open()
    for i in range(300):
        page = doc.new_page()
        page.insert_text((72, 72), f"Page {i + 1} body text for median benchmark.", fontsize=11)
        if i % 50 == 0:
            page.insert_text((72, 50), f"SECTION {i // 50 + 1}", fontsize=16)
    data = doc.tobytes()
    doc.close()

    t0 = time.perf_counter()
    result = run_ocr_cleaning(file_type=FileType.pdf, data=data)
    elapsed = time.perf_counter() - t0

    assert result.page_count == 300
    assert result.char_count > 0
    assert elapsed < 30.0  # generous CI bound; sampling keeps median cheap
